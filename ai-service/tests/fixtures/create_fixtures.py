"""
Generate synthetic test resume fixtures with known content.
Run once to create PDF and DOCX files, then commit them.

Usage: python -m tests.fixtures.create_fixtures
"""
import os

FIXTURES_DIR = os.path.dirname(os.path.abspath(__file__))

RESUME_TEXT = """RAHUL SHARMA
rahul.sharma@example.com | +91-9876543210 | Mumbai, Maharashtra, India

SUMMARY
Experienced software developer with 3 years of expertise in full-stack web development.

EDUCATION
B.Tech in Computer Science
Indian Institute of Technology, Mumbai | 2020 | GPA: 8.5/10

EXPERIENCE
Software Developer | Tech Solutions Pvt Ltd | Jan 2021 - Jan 2023
- Developed web applications using React and Node.js
- Built RESTful APIs serving 10,000+ daily requests

Junior Developer | StartupXYZ | Jan 2020 - Dec 2020
- Built REST APIs using Flask and Python
- Implemented database schemas in MongoDB

SKILLS
Technical: Python, JavaScript, React, Node.js, MongoDB, Flask
Soft: Communication, Team Leadership, Problem Solving

PROJECTS
E-Commerce Platform
- Full-stack e-commerce with React and Node.js
- Technologies: React, Node.js, MongoDB

CERTIFICATIONS
- AWS Certified Developer
"""


def create_pdf():
    """Create a simple PDF resume using reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas

        path = os.path.join(FIXTURES_DIR, "sample_resume.pdf")
        c = canvas.Canvas(path, pagesize=A4)
        width, height = A4
        y = height - 50

        for line in RESUME_TEXT.strip().split("\n"):
            if y < 50:
                c.showPage()
                y = height - 50
            c.drawString(50, y, line)
            y -= 15

        c.save()
        print(f"Created: {path}")
    except ImportError:
        print("reportlab not installed, creating minimal PDF fallback")
        path = os.path.join(FIXTURES_DIR, "sample_resume.pdf")
        lines = RESUME_TEXT.strip().replace("\n", " ")
        content = f"""%PDF-1.4
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj
4 0 obj<</Length {len(lines) + 30}>>stream
BT /F1 10 Tf 50 750 Td ({lines})Tj ET
endstream endobj
5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj
xref
0 6
trailer<</Size 6/Root 1 0 R>>
startxref
0
%%EOF"""
        with open(path, "w") as f:
            f.write(content)
        print(f"Created (basic fallback): {path}")


def create_docx():
    """Create a simple DOCX resume using python-docx."""
    from docx import Document

    path = os.path.join(FIXTURES_DIR, "sample_resume.docx")
    doc = Document()

    for line in RESUME_TEXT.strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        if line == line.upper() and len(line) > 3 and not line.startswith("-"):
            doc.add_heading(line, level=2)
        else:
            doc.add_paragraph(line)

    doc.save(path)
    print(f"Created: {path}")


if __name__ == "__main__":
    create_pdf()
    create_docx()
    print("Done! Fixtures created successfully.")
